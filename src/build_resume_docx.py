#!/usr/bin/env python3
"""从 resume.md 生成 docx 简历文件（保留加粗与结构）"""
import re
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

SRC = "/Users/wenbowang/Documents/trae_projects/knowledge-sharing/interview-prep/resume.md"
OUT = "/Users/wenbowang/Documents/trae_projects/knowledge-sharing/interview-prep/王文博_AI产品经理简历_v2.docx"

doc = Document()
# 页边距
for sec in doc.sections:
    sec.top_margin = Cm(1.5)
    sec.bottom_margin = Cm(1.5)
    sec.left_margin = Cm(1.8)
    sec.right_margin = Cm(1.8)

def set_font(run, size=10.5, bold=False, color=None):
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.name = "微软雅黑"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    if color:
        run.font.color.rgb = color

def add_para(text, size=10.5, bold=False, align=None, space_after=4):
    p = doc.add_paragraph()
    if align:
        p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    # 粗体部分用 ** 标记
    parts = re.split(r"(\*\*.*?\*\*)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            r = p.add_run(part[2:-2])
            set_font(r, size, bold=True)
        else:
            r = p.add_run(part)
            set_font(r, size, bold=bold)
    return p

lines = open(SRC, encoding="utf-8").read().splitlines()

# 跳过文件头注释（## 之前的 --- 块）
i = 0
while i < len(lines):
    line = lines[i].strip()
    if not line:
        i += 1
        continue
    if line == "---":
        # 跳过块引用注释
        i += 1
        while i < len(lines) and (lines[i].strip().startswith(">") or not lines[i].strip()):
            i += 1
        continue
    if line.startswith(">"):
        i += 1
        continue
    if line.startswith("# 王文博"):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(6)
        r = p.add_run("王文博")
        set_font(r, 20, bold=True)
        i += 1
        continue
    if line.startswith("#"):
        title = line.lstrip("#").strip()
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(4)
        r = p.add_run(title)
        set_font(r, 13, bold=True, color=RGBColor(0x1F, 0x3B, 0x73))
        i += 1
        continue
    if line.startswith("### "):
        # 实习/项目条目标题：提取加粗部分与右侧信息
        text = line.lstrip("#").strip()
        add_para(text, size=11, bold=True, space_after=3)
        i += 1
        continue
    if line.startswith("- "):
        add_para(line[2:], space_after=3)
        i += 1
        continue
    # 普通段落
    add_para(line)
    i += 1

doc.save(OUT)
print(f"saved: {OUT}")
